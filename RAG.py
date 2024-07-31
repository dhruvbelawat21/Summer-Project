
!pip install faiss-cpu faiss-gpu ujson python-docx doc2docx openai gitpython
import os
import sys
import random
import numpy as np
import pandas as pd
import torch
from string import Template
import json
from tqdm.auto import tqdm
import warnings
from transformers import pipeline

warnings.simplefilter("ignore")
from transformers import AutoModelForCausalLM, AutoTokenizer
from doc2docx import convert
from docx import Document
import sqlite3
import faiss
import ujson
from copy import deepcopy
import traceback
import time
from openai import OpenAI
import numpy as np
import json
import git
from tqdm.auto import tqdm
import chardet
import ast
import openai
import re
import torch
from torch.utils.data import DataLoader, TensorDataset
import torch.nn as nn
import torch.nn.functional as F
import logging
from google.colab import userdata
api_key = userdata.get('OpenAI')
folder_url = "https://huggingface.co/datasets/netop/Embeddings3GPP-R18"
clone_directory = "./3GPP-Release18"

if not os.path.exists(clone_directory):
    git.Repo.clone_from(folder_url, clone_directory)
    print("Folder cloned successfully!")
else:
    print("Folder already exists. Skipping cloning.")
device = "cuda" if torch.cuda.is_available() else "cpu"
torch.set_default_device(device)
model = AutoModelForCausalLM.from_pretrained("microsoft/phi-2", torch_dtype="auto", trust_remote_code=True)
tokenizer = AutoTokenizer.from_pretrained("microsoft/phi-2", trust_remote_code=True)
from transformers import AutoTokenizer, AutoModelForCausalLM

# Initialize the tokenizer and model for phi-2
tokenizer = AutoTokenizer.from_pretrained("microsoft/phi-2")
tokenizer.pad_token = tokenizer.eos_token

model = AutoModelForCausalLM.from_pretrained("microsoft/phi-2")

def phi2(prompt):
    # Tokenize the input prompt
    inputs = tokenizer(prompt, return_tensors="pt", return_attention_mask=True, padding=True)

    # Generate text
    outputs = model.generate(input_ids=inputs['input_ids'], attention_mask=inputs['attention_mask'], max_length=2048)

    # Decode the generated text
    text = tokenizer.batch_decode(outputs)[0]

    return text
'''
#prompt = """
        is a 2-kg feather heavier than a 1-kg feather
"""

#phi2(prompt)
'''
class Storage:
    def __init__(self, db_name):
        try:
            self.db_name = db_name
            self.conn = sqlite3.connect(db_name, check_same_thread=False)  # This allows the connection to be used in multiple threads
            self.cursor = self.conn.cursor()
            self.optimize_database()
        except sqlite3.Error as e:
            print(f"Failed to connect to database {db_name}: {e}")
            raise e  # Re-raise exception after logging to handle it upstream if needed

    def __enter__(self):
        self.conn = sqlite3.connect(self.db_name, check_same_thread=False)
        self.cursor = self.conn.cursor()
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.close()

    def optimize_database(self):
        self.cursor.execute("PRAGMA cache_size = -64000;")
        self.cursor.execute("PRAGMA journal_mode = WAL;")
        self.cursor.execute("PRAGMA synchronous = OFF;")

    def create_dataset(self, dataset_name):
        try:
            # Create a new table with 'id' as the primary key and 'data' for JSON storage, if it doesn't already exist.
            self.cursor.execute(f'''CREATE TABLE IF NOT EXISTS {dataset_name} (id TEXT PRIMARY KEY, data TEXT)''')
            self.conn.commit()  # Commit changes to the database.
        except sqlite3.Error as e:
            # Handle SQLite errors, e.g., syntax errors in the SQL command.
            print(dataset_name)
            print(f"Database error: {e}")
        except Exception as e:
            # Handle unexpected errors, keeping the program from crashing.
            print(f"Exception in create_dataset: {e}")


    def insert_dict(self, dataset_name, data_dict):
        # Extract 'id' from data dictionary.
        dict_id = data_dict.get('id')

        # Proceed only if 'id' exists and it's not already in the dataset.
        if dict_id is not None and not self.is_id_in_dataset(dataset_name, dict_id):
            try:
                # Convert dictionary to JSON string.
                data_json = ujson.dumps(data_dict)
                # Insert 'id' and JSON string into the specified dataset.
                self.cursor.execute(f"INSERT INTO {dataset_name} (id, data) VALUES (?, ?)", (dict_id, data_json,))
                self.conn.commit()  # Commit the transaction.
            except sqlite3.Error as e:
                # Handle SQLite errors during insert operation.
                print(f"Database error: {e}")
            except Exception as e:
                # Handle other unexpected errors.
                print(f"Exception in insert_dict: {e}")

    def insert_dict_new(self, dataset_name, document, index):
        import json
        # Insert a dictionary as a JSON string into the specified table
        try:
            data_str = json.dumps(document)
            # Check if the ID already exists in the table
            self.cursor.execute(f"SELECT id FROM {dataset_name} WHERE id = ?", (index,))
            existing_id = self.cursor.fetchone()
            if existing_id:
                print(f"ID {index} already exists in {dataset_name}. Clearing table and retrying insertion.")
                # Clear the table
                self.cursor.execute(f"DELETE FROM {dataset_name}")
                self.conn.commit()
                print(f"All existing data in {dataset_name} has been cleared.")
            # Insert the new data
            self.cursor.execute(f"INSERT INTO {dataset_name} (id, data) VALUES (?, ?)", (index, data_str))
            self.conn.commit()
            print(f"Document inserted successfully into {dataset_name} with ID {index}.")
        except Exception as e:
            print(f"Error inserting into table {dataset_name}: {e}")


    def insert_or_update_dict(self, dataset_name, data_dict):
        # Extract 'id' from data dictionary.
        dict_id = data_dict.get('id')

        # Proceed only if 'id' exists.
        if dict_id is not None:
            try:
                # Convert dictionary to JSON string.
                data_json = ujson.dumps(data_dict)
                # Use 'REPLACE INTO' to insert or update the row with the specified 'id'.
                self.cursor.execute(f"REPLACE INTO {dataset_name} (id, data) VALUES (?, ?)", (dict_id, data_json,))
                self.conn.commit()  # Commit the transaction.
            except sqlite3.Error as e:
                # Handle SQLite errors during insert/update operation.
                print(f"Database error: {e}")
            except Exception as e:
                # Handle other unexpected errors.
                print(f"Exception in insert_or_update_dict: {e}")
                print(traceback.format_exc())


    def is_id_in_dataset(self, dataset_name, dict_id):
        try:
            # Execute a SQL query to check if the given 'dict_id' exists in the 'dataset_name' table.
            self.cursor.execute(f"SELECT 1 FROM {dataset_name} WHERE id = ?", (dict_id,))
            # Return True if the ID exists, False otherwise.
            return self.cursor.fetchone() is not None
        except sqlite3.Error as e:
            # Handle SQLite errors, logging the issue and indicating the ID was not found.
            print(f"Database error: {e}")
            return False


    def store_faiss_data(self, identifier, index, data_mapping):
        try:
            # Serialize the FAISS index into bytes for storage.
            serialized_index = faiss.serialize_index(index).tobytes()
            # Convert the data mapping dictionary into a JSON string.
            json_data_mapping = ujson.dumps(data_mapping)

            # Ensure the storage table exists, creating it if necessary.
            self.cursor.execute('''CREATE TABLE IF NOT EXISTS faiss_index_data
                                (id TEXT PRIMARY KEY,
                                    faiss_index BLOB,
                                    data_mapping TEXT)''')

            # Insert the serialized index and JSON data mapping into the database.
            self.cursor.execute("INSERT INTO faiss_index_data (id, faiss_index, data_mapping) VALUES (?, ?, ?)",
                                (identifier, serialized_index, json_data_mapping,))
            self.conn.commit()  # Commit the changes to the database.
        except sqlite3.Error as e:
            # Log database-related errors.
            print(f"Database error: {e}")
        except Exception as e:
            # Log any other exceptions.
            print(f"Exception in store_faiss_data: {e}")


    def retrieve_faiss_data(self, identifier):
        try:
            # Correct the SQL query to fetch the faiss_index and data_mapping for the given identifier.
            # Remove the parentheses around the selected columns to ensure proper data retrieval.
            self.cursor.execute("SELECT faiss_index, data_mapping FROM faiss_index_data WHERE id = ?", (identifier,))
            row = self.cursor.fetchone()

            if row is not None:
                # Correctly deserialize the FAISS index from the binary data stored in the database.
                # Use faiss.deserialize_index directly on the binary data without converting it back to an array.
                index = faiss.deserialize_index(row[0])

                # Deserialize the JSON string back into a Python dictionary.
                data_mapping = ujson.loads(row[1])

                return index, data_mapping
            else:
                # Return None if no entry was found for the given identifier.
                return None, None
        except sqlite3.Error as e:
            # Log SQLite errors and return None to indicate failure.
            print(f"Database error: {e}")
            return None, None
        except Exception as e:
            # Log unexpected errors and return None to indicate failure.
            print(f"Exception in retrieve_faiss_data: {e}")

    def retrieve_dicts(self, dataset_name):
        # Properly quote the table name to prevent SQL injection and syntax errors.
        safe_dataset_name = f'"{dataset_name}"'
        start= time.time()
        # Execute a SQL query to fetch all records from the specified dataset table.
        self.cursor.execute(f"SELECT * FROM {safe_dataset_name}")
        rows = self.cursor.fetchall()
        end=time.time()
        print(f"-------------------------{end-start}")
        # Utilizing a faster JSON parsing library if significant JSON parsing overhead is detected
        start= time.time()
        a= [ujson.loads(row[1]) for row in rows]
        end=time.time()
        print(f"UJSON-------------------------{end-start}")

        return a
    device = "cuda:0" if torch.cuda.is_available() else "cpu"

# Load test data
    with open("/content/TeleQnA_testing1.txt", "r") as f:
        test_data = f.read()
    test_data = json.loads(test_data)

    # Function to generate responses and store in context.txt
    def generate_responses():
        # Initialize pipeline
        phi3_telecom = pipeline("text-generation", model="EleutherAI/gpt-neo-125M")

        # Dictionary to hold question IDs and responses
        context_dict = {}

        for question_id, question_details in tqdm(test_data.items()):
            original_query = question_details["question"]
            response = phi3_telecom(original_query, max_new_tokens=50)[0]['generated_text'].strip()
            context_dict[question_id] = response

        # Write dictionary to context.txt in JSON format
        with open("context.txt", "w") as file:
            json.dump(context_dict, file, indent=4)

    # Generate responses and store in context.txt
    generate_responses()

    def reset_database(self):
        try:
            # Retrieve the names of all tables in the database.
            self.cursor.execute("SELECT name FROM sqlite_master WHERE type='table'")
            tables = self.cursor.fetchall()

            # Iterate over each table name and drop the table to remove it from the database.
            for table in tables:
                self.cursor.execute(f"DROP TABLE {table[0]}")

            # Commit the changes to finalize the removal of all tables.
            self.conn.commit()
        except sqlite3.Error as e:
            # Handle and log any SQLite errors encountered during the operation.
            print(f"Database error: {e}")
        except Exception as e:
            # Handle and log any non-SQLite errors that may occur.
            print(f"Exception in reset_database: {e}")

    def get_dict_by_id(self, dataset_name, dict_id):
        try:
            # Execute SQL query to fetch the record with the specified ID from the given dataset.
            self.cursor.execute(f"SELECT data FROM {dataset_name} WHERE id = ?", (dict_id,))
            row = self.cursor.fetchone()

            # If the record exists, convert the JSON string in 'data' column back to a dictionary and return it.
            if row is not None:
                return ujson.loads(row[0])
            else:
                # Return None if no record was found with the given ID.
                return None
        except sqlite3.Error as e:
            # Log database-related errors and return None to indicate failure.
            print(f"Database error: {e}")
            Storage.get_dict_by_id(self, dataset_name, dict_id)

        except Exception as e:
            # Log any other exceptions that occur and return None to indicate failure.
            print(f"Exception in get_dict_by_id: {e}")

    def close(self):
        try:
            # Attempt to close the database connection.
            self.conn.close()
        except sqlite3.Error as e:
            # Log any SQLite errors encountered during the closing process.
            print(f"Database error: {e}")
        except Exception as e:
            # Log any other exceptions that might occur during closure.
            print(f"Exception in close: {e}")
            print(traceback.format_exc())
import logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def read_docx(file_path):
    """Read and extract text from a DOCX file."""
    try:
        doc = Document(file_path)
        return '\n'.join(para.text for para in doc.paragraphs)
    except Exception as e:
        logging.error(f"Failed to read DOCX file at {file_path}: {e}")
        return None

def get_documents(series_list, folder_path=r'/content/3GPP-Release18/Documents', storage_name='Documents.db', dataset_name="Standard"):
    """Retrieve and process documents from a folder, storing them in a database if not already present."""
    storage = Storage(f'./3GPP-Release18/{storage_name}')
    storage.create_dataset(dataset_name)

    document_ds = []
    file_list = []

    # Check and convert .doc files to .docx
    convert_docs_to_docx(folder_path)

    # Prepare list of .docx files for processing
    file_list = [f for f in tqdm(os.listdir(folder_path), desc="Filtering documents") if valid_file(f, series_list)]

    # Process each document
    for filename in tqdm(file_list, desc="Processing documents"):
        file_path = os.path.join(folder_path, filename)
        process_document(file_path, filename, storage, document_ds, dataset_name)

    storage.close()
    return document_ds

def convert_docs_to_docx(folder_path):
    """Convert .doc files in a folder to .docx format if any."""
    has_doc = any(f.endswith('.doc') for f in os.listdir(folder_path))
    if has_doc:
        convert(folder_path)

def valid_file(filename, series_list):
    """Check if a file should be processed based on its name and series list."""
    return filename.endswith(".docx") and not filename.startswith("~$") and (not filename[:2].isnumeric() or int(filename[:2]) in series_list)

def process_document(file_path, filename, storage, document_ds, dataset_name):
    """Process a single document file."""
    if storage.is_id_in_dataset(dataset_name, filename):
        data_dict = storage.get_dict_by_id(dataset_name, filename)
        document_ds.append(data_dict)
    else:
        content = read_docx(file_path)

        if content:
            data_dict = {'id': filename, 'text': content, 'source': filename}
            document_ds.append(data_dict)
            storage.insert_dict(dataset_name, data_dict)
def custom_text_splitter(text, chunk_size, chunk_overlap, word_split=False):
    """
    Splits a given text into chunks of a specified size with a defined overlap between them.

    This function divides the input text into chunks based on the specified chunk size and overlap.
    Optionally, it can split the text at word boundaries to avoid breaking words when 'word_split'
    is set to True. This is achieved by using a regular expression that identifies word separators.

    Args:
        text (str): The text to be split into chunks.
        chunk_size (int): The size of each chunk in characters.
        chunk_overlap (int): The number of characters of overlap between consecutive chunks.
        word_split (bool, optional): If True, ensures that chunks end at word boundaries. Defaults to False.

    Returns:
        list of str: A list containing the text chunks.
    """
    chunks = []
    start = 0
    separators_pattern = re.compile(r'[\s,.\-!?{}":;<>]+')

    while start < len(text) - chunk_overlap:
        end = min(start + chunk_size, len(text))

        if word_split:
            match = separators_pattern.search(text, end)
            if match:
                end = match.end()

        if end == start:
            end = start + 1

        chunks.append(text[start:end])
        start = end - chunk_overlap

        if word_split:
            match = separators_pattern.search(text, start-1)
            if match:
                start = match.start() + 1

        if start < 0:
            start = 0

    return chunks


def chunk_doc(doc):
    chunks= custom_text_splitter( doc["text"], 500, 25, word_split = True)
    return [{"text": chunk, "source": doc["source"]} for chunk in chunks]
def get_embeddings(series_docs):
    """Add embeddings to each chunk of documents from pre-saved NumPy files."""
    for doc_key, doc_chunks in series_docs.items():
        try:
            # Load embeddings specific to each document series
            embeddings = np.load(f'3GPP-Release18\Embeddings\Embeddings{doc_key}.npy')
            # print(f'{len(embeddings)}---{doc_key}')

            # print(f"{len(np.unique([str(a[0])+str(a[1]) for a in embeddings]))} -- {len(embeddings)}" )
        except FileNotFoundError:
            logging.error(f"Embedding file for {doc_key} not found.")
            continue
        except Exception as e:
            logging.error(f"Failed to load embeddings for {doc_key}: {e}")
            continue

        text_list=[]
        for chunk in doc_chunks:
            for single_chunk in chunk:
                text_list.append(single_chunk['text'])
        dex ={}
        for i in range(len(text_list)):
            dex[text_list[i]] = embeddings[i]
        # Process each chunk within the document series
        updated_chunks = []
        for chunk in doc_chunks:
            for idx, single_chunk in enumerate(chunk):
                try:
                    chunk[idx]['embedding'] = dex[chunk[idx]['text']]
                    updated_chunks.append(chunk[idx])
                except IndexError:
                    logging.warning(f"Embedding index {idx} out of range for {doc_key}.")
                except Exception as e:
                    logging.error(f"Error processing chunk {idx} for {doc_key}: {e}")

        series_docs[doc_key] = updated_chunks
    return series_docs
!pip install python-docx
from docx import Document




def read_docx(file_path):
    #Reads a .docx file and categorizes its content into terms and abbreviations.
    doc = Document(file_path)

    processing_terms = False
    processing_abbreviations = False
    start = 0
    terms_definitions = {}
    abbreviations_definitions = {}

    for para in doc.paragraphs:
        text = para.text.strip()
        if "References" in text:
            start += 1
        if start >=2:
            if "Terms and definitions" in text:
                processing_terms = True
                processing_abbreviations = False

            elif "Abbreviations" in text:
                processing_abbreviations = True
                processing_terms = False
            else:
                if processing_terms and ':' in text:
                    term, definition = text.split(':', 1)
                    terms_definitions[term.strip()] = definition.strip().rstrip('.')
                elif processing_abbreviations and '\t' in text:
                    abbreviation, definition = text.split('\t', 1)
                    if len(abbreviation)> 1:
                        abbreviations_definitions[abbreviation.strip()] = definition.strip()

    return terms_definitions, abbreviations_definitions

file_path = "/content/3GPP-Release18/Documents/21900-i10.docx"
terms_definitions, abbreviations_definitions = read_docx(file_path)

def preprocess(text, lowercase=True):
    #Converts text and optionally converts to lowercase. Removes punctuation.
    if lowercase:
        text = text.lower()
    punctuations = "!()-[]{};:'\,<>./?@#$%^&*_~"
    for char in punctuations:
        text = text.replace(char, '')
    return text

def find_and_filter_terms(terms_dict, sentence):
    #Finds terms in the given sentence, case-insensitively, and filters out shorter overlapping terms.
    lowercase_sentence = preprocess(sentence, lowercase=True)

    # Find all terms
    matched_terms = {term: terms_dict[term] for term in terms_dict if preprocess(term) in lowercase_sentence}

    # Filter out terms that are subsets of longer terms
    final_terms = {}
    for term in matched_terms:
        if not any(term in other and term != other for other in matched_terms):
            final_terms[term] = matched_terms[term]

    return final_terms

def find_and_filter_abbreviations(abbreviations_dict, sentence):
    """Finds abbreviations in the given sentence, case-sensitively, and filters out shorter overlapping abbreviations."""
    processed_sentence = preprocess(sentence, lowercase=False)
    words = processed_sentence.split()

    matched_abbreviations = {word: abbreviations_dict[word] for word in words if word in abbreviations_dict}

    final_abbreviations = {}
    sorted_abbrs = sorted(matched_abbreviations, key=len, reverse=True)
    for abbr in sorted_abbrs:
        if not any(abbr in other and abbr != other for other in sorted_abbrs):
            final_abbreviations[abbr] = matched_abbreviations[abbr]

    print(final_abbreviations)
    return final_abbreviations

def find_terms_and_abbreviations_in_sentence(terms_dict, abbreviations_dict, sentence):
    """Finds, filters terms and abbreviations in the given sentence.
       Filters to prioritize longer terms and abbreviations."""
    matched_terms = find_and_filter_terms(terms_dict, sentence)
    matched_abbreviations = find_and_filter_abbreviations(abbreviations_dict, sentence)

    # Format matched terms and abbreviations for output
    formatted_terms = [f"{term}: {definition}" for term, definition in matched_terms.items()]
    formatted_abbreviations = [f"{abbr}: {definition}" for abbr, definition in matched_abbreviations.items()]

    return formatted_terms, formatted_abbreviations

def find_terms_and_abbreviations_in_sentence(terms_dict, abbreviations_dict, sentence):
    """Finds and filters terms or abbreviations in the given sentence.
       Abbreviations are matched case-sensitively, terms case-insensitively, and longer terms are prioritized."""
    processed_sentence = preprocess(sentence, lowercase=False)  # Preserve case for abbreviations
    matched_abbreviations = {abbr: abbreviations_dict[abbr] for abbr in abbreviations_dict if abbr in processed_sentence}

    # Find and filter terms
    matched_terms = find_and_filter_terms(terms_dict, sentence)

    # Format matched terms and abbreviations for output
    formatted_terms = [f"{term}: {definition}" for term, definition in matched_terms.items()]
    formatted_abbreviations = [f"{abbr}: {definition}" for abbr, definition in matched_abbreviations.items()]

    return formatted_terms, formatted_abbreviations

def get_def(sentence):
    formatted_terms, formatted_abbreviations = find_terms_and_abbreviations_in_sentence(terms_definitions, abbreviations_definitions, sentence)
    defined = []
    for term in formatted_terms:
        defined.append(term[:3])
    for abbreviation in formatted_abbreviations:
        defined.append(abbreviation[:3])


def define_TA_question(sentence):
    formatted_terms, formatted_abbreviations = find_terms_and_abbreviations_in_sentence(terms_definitions, abbreviations_definitions, sentence)
    terms= '\n'.join(formatted_terms)
    abbreviations= '\n'.join(formatted_abbreviations)
    question= f"""{sentence}\n
Terms and Definitions:\n
{terms}\n

Abbreviations:\n
{abbreviations}\n
"""
    return question
import os
os.environ['OMP_NUM_THREADS'] = '8'

def create_faiss_index_IndexFlatIP(embeddings, data, source):
    """Create FAISS IndexFlatIP from embeddings and maps indices to data and source."""
    try:
        logging.info("Creating IndexFlatIP...")
        d = embeddings.shape[1]
        index = faiss.IndexFlatIP(d)
        index.add(embeddings)
        index_to_data_mapping = {i: data[i] for i in range(len(data))}
        index_to_source_mapping = {i: source[i] for i in range(len(source))}
        return index, index_to_data_mapping, index_to_source_mapping
    except Exception as e:
        logging.error(f"Error creating FAISS index: {e}")
        return None, None, None

def get_faiss_batch_index(embedded_batch):
    """Generate FAISS index from a batch of embeddings, handling missing embeddings by generating them."""
    try:
        source = [chunk['source'] for chunked_batch in embedded_batch for chunk in chunked_batch]
        embeddings = []
        data = []

        for doc in embedded_batch:
            embeddings_batch = []
            for chunk in doc:
                if 'embedding' in chunk:
                    embeddings_batch.append(chunk['embedding'])
                else:
                    embedding = generate_embedding_for_chunk(chunk)
                    if embedding is not None:
                        chunk['embedding'] = embedding
                        embeddings_batch.append(embedding)

            embeddings.extend(embeddings_batch)
            data.extend([chunk['text'] for chunk in doc])

        embeddings = np.array(embeddings, dtype=np.float32)
        return create_faiss_index_IndexFlatIP(embeddings, data, source)
    except Exception as e:
        logging.error(f"Failed to process batch for FAISS indexing: {e}")
        return None, None, None

def generate_embedding_for_chunk(chunk):
    """Generate embeddings for a chunk using the OpenAI API."""
    try:
        client = OpenAI()
        response = client.embeddings.create(
            input=chunk["text"],
            model="text-embedding-3-large"
        )
        return response['data'][0]['embedding']
    except Exception as e:
        logging.error(f"Failed to generate embedding for chunk: {chunk['text']}. Error: {e}")
        return None
def search_faiss_index(faiss_index, query_embedding, k=5):
    # Validate input parameters
    if not isinstance(query_embedding, np.ndarray) or query_embedding.ndim != 1:
        raise ValueError("query_embedding must be a 1D numpy array")
    if not isinstance(k, int) or k <= 0:
        raise ValueError("k must be a positive integer")

    # Reshape the query embedding for FAISS (FAISS expects a 2D array)
    query_embedding_reshaped = query_embedding.reshape(1, -1)

    # Perform the search
    D, I = faiss_index.search(query_embedding_reshaped, k)

    # Return the indices and distances of the nearest neighbors
    return I, D

def get_query_embedding_OpenAILarge(query_text, context=None):
    try:
        if context is not None:
            query_text = f'{query_text}\n' + "\n".join(context)

        client = OpenAI()
        response = client.embeddings.create(
            input=query_text,
            model="text-embedding-3-large",
            dimensions= 1024
            )
        query_embedding =  response.data[0].embedding
        return np.array(query_embedding, dtype=np.float32)
    except Exception as e:
        print(f"Error occurred in get_query_embedding_OpenAILarge: {e}")
        traceback.print_exc()  # Using `print_exc` for consistency

def find_nearest_neighbors_faiss(query_text, faiss_index, data_mapping, k, source_mapping,  context=None):

    try:
        query_embedding = get_query_embedding_OpenAILarge(query_text, context)

        I, D = search_faiss_index(faiss_index, query_embedding, k)

        nearest_neighbors = []
        for index in I[0]:
            if index < len(data_mapping):
                data = data_mapping.get(index, "Data not found")
                source = source_mapping.get(index, "Source not found")
                nearest_neighbors.append((index, data, source))
        return nearest_neighbors
    except Exception as e:
        print(f"Error in find_nearest_neighbors_faiss: {str(e)}")
        traceback.print_exc()
        return []
import torch
import torch.nn as nn
import torch.nn.functional as F

class NNRouter(nn.Module):
    def __init__(self):
        super(NNRouter, self).__init__()
        self.layer1_1 = nn.Linear(1024, 768)
        self.layer1_2 = nn.Linear(768, 512)
        self.layer1_3 = nn.Linear(512, 256)
        self.dropout1 = nn.Dropout(0.2)

        self.layer2_1 = nn.Linear(18, 128)
        self.layer2_2 = nn.Linear(128, 256)
        self.dropout2 = nn.Dropout(0.05)

        self.batchnorm1 = nn.BatchNorm1d(256)
        self.batchnorm2 = nn.BatchNorm1d(256)

        self.alfa = nn.Parameter(torch.ones(1), requires_grad=True)
        self.beta = nn.Parameter(torch.ones(1), requires_grad=True)

        self.output_layer1 = nn.Linear(256, 128)
        self.output_layer2 = nn.Linear(128, 18)

        self.leaky_relu = nn.LeakyReLU(0.01)

    def forward(self, input_1, input_2):

        x1 = F.relu(self.layer1_1(input_1))
        x1 = self.dropout1(x1)
        x1 = F.relu(self.layer1_2(x1))
        x1 = self.dropout1(x1)
        x1 = F.relu(self.layer1_3(x1))
        x1 = self.batchnorm1(x1)

        x2 = F.relu(self.layer2_1(input_2))
        x2 = self.dropout2(x2)
        x2 = F.relu(self.layer2_2(x2))
        x2 = self.batchnorm2(x2)

        weighted_x1 = self.alfa * x1
        weighted_x2 = self.beta * x2

        combined = weighted_x1 + weighted_x2

        output = self.output_layer1(self.leaky_relu(combined))
        output = self.output_layer2(self.leaky_relu(output))

        return output
import zipfile
from google.colab import userdata
import openai
# from openai import OpenAI
api_key = userdata.get('OpenAI')
class Query:
    def __init__(self, query, context):
        self.id = id
        self.question = query
        self.query = query
        self.enhanced_query = query
        self.con_counter = {}
        self.topic_distr = []
        if isinstance(context, str):
            context = [context]
        self.context = context
        self.rowcontext = []
        self.context_source = []
        self.possible_sources = []
        self.wg = []
        self.source_hit={}
        self.document_accuracy = None


    def def_TA_question(self):
        self.query = define_TA_question(self.query)
        self.enhanced_query = self.query

    def candidate_answers(self):
            try:
                client = OpenAI()
                try:
                    row_context = f"""
                    Provide all the possible answers to the fallowing question. Conisdering your knowledge and the text provided.
                    Question {self.query}\n

                    Considering the fallowing context:
                    {self.context}

                    Provide all the possible answers to the fallowing question. Conisdering your knowledge and the text provided.
                    Question {self.question}\n


                    Make sure none of the answers provided contradicts with your knowledge and have at most 100 characters each.
                    """
                    generated_output = client.chat.completions.create(
                    model = "gpt-3.5-turbo-1106",
                    messages = [
                        {"role": "system", "content": "You are an expert at telecom knowledge. Be concise, precise and provide exact technical terms required."},
                        {"role": "user", "content":  row_context},

                    ],
                    )
                    generated_output_str = generated_output.choices[0].message.content
                    print(generated_output_str)
                    if generated_output_str != "NO":
                        self.context = generated_output_str
                        self.enhanced_query = self.query +'\n'+ self.context
                except Exception as e:
                    print(f"An error occurred: {e}")
            except:
                print("ERROR")
                print(traceback.format_exc())

    def get_embeddings_list(text_list):
        # Initialize the OpenAI client
        api_key = userdata.get('OpenAI')
        client = openai.OpenAI(api_key=api_key)
        # Request embeddings for the list of texts using a different, larger model
        response = client.embeddings.create(
                    input=text_list,
                    model="text-embedding-3-large",
                    dimensions=1024,
                )
        embeddings = []
        for i in range(len(response.data)):
            embeddings.append(response.data[i].embedding)

        text_embeddings= {}
        print(len(text_list))
        print(len(embeddings))
        for index in range(len(text_list)):
            text_embeddings[text_list[index]] = embeddings[index]
        return text_embeddings

    def inner_product(a, b):
        """Compute the inner product of two lists."""
        return sum(x * y for x, y in zip(a, b))

    def get_col2(embeddings_list):
        topics_with_series = [("Requirements (21 series): Focuses on the overarching requirements necessary for UMTS (Universal Mobile Telecommunications System) and later cellular standards, including GSM enhancements, security standards, and the general evolution of 3GPP systems. It covers vocabulary, security threats, UE capability requirements, and work items for various releases.", "21 series"),
            ("Service aspects ('stage 1') (22 series): This series details the initial specifications for services provided by the network, outlining the service requirements before the technical realization is detailed. It serves as the first step in defining what the network should provide.", "22 series"),
            ("Technical realization ('stage 2') (23 series): Focuses on the architectural and functional framework necessary to implement the services described in stage 1, providing a bridge to the detailed protocols and interfaces defined in stage 3​.,", "23 series"),
            ("Signalling protocols ('stage 3') - user equipment to network (24 series): Details the protocols and signaling procedures for communication between user equipment and the network, ensuring interoperability and successful service delivery.", "24 series"),
            ("Radio aspects (25 series): Covers the specifications related to radio transmission technologies, including frequency bands, modulation schemes, and antenna specifications, critical for ensuring efficient and effective wireless communication​.", "25 series"),
            ("CODECs (26 series): Contains specifications for voice, audio, and video codecs used in the network, defining how data is compressed and decompressed to enable efficient transmission over bandwidth-limited wireless networks.", "26 series"),
            ("Data (27 series): This series focuses on the data services and capabilities of the network, including specifications for data transmission rates, data service features, and support for various data applications.", "27 series"),
            ("Signalling protocols ('stage 3') - (RSS-CN) and OAM&P and Charging (overflow from 32.- range) (28 series): Addresses additional signaling protocols related to operation, administration, maintenance, provisioning, and charging, complementing the core signaling protocols outlined in the 24 series.", "28 series"),
            ("Signalling protocols ('stage 3') - intra-fixed-network (29 series): Specifies signaling protocols used within the fixed parts of the network, ensuring that various network elements can communicate effectively to provide seamless service to users.", "29 series"),
            ("Programme management (30 series): Relates to the management and coordination of 3GPP projects and work items, including documentation and specification management procedures​.", "30 series"),
            ("Subscriber Identity Module (SIM / USIM), IC Cards. Test specs. (31 series): Covers specifications for SIM and USIM cards, including physical characteristics, security features, and interaction with mobile devices, as well as testing specifications for these components​.", "31 series"),
            ("OAM&P and Charging (32 series): Focuses on operation, administration, maintenance, and provisioning aspects of the network, as well as the charging principles and mechanisms for billing and accounting of network services.", "32 series"),
            ("Security aspects (33 series): Details the security mechanisms and protocols necessary to protect network operations, user data, and communication privacy, including authentication, encryption, and integrity protection measures​.", "33 series"),
            ("UE and (U)SIM test specifications (34 series): Contains test specifications for User Equipment (UE) and (U)SIM cards, ensuring that devices and SIM cards meet 3GPP standards and perform correctly in the network​.", "34 series"),
            ("Security algorithms (35 series): Specifies the cryptographic algorithms used in the network for securing user data and signaling information, including encryption algorithms and key management procedures.", "35 series"),
            ("LTE (Evolved UTRA), LTE-Advanced, LTE-Advanced Pro radio technology (36 series): Details the technical specifications for LTE, LTE-Advanced, and LTE-Advanced Pro technologies, including radio access network (RAN) protocols, modulation schemes, and network architecture​.", "36 series"),
            ("Multiple radio access technology aspects (37 series): Addresses the integration and interoperability of multiple radio access technologies within the network, enabling seamless service across different types of network infrastructure.", "37 series"),
            ("Radio technology beyond LTE (38 series): Focuses on the development and specification of radio technologies that extend beyond the capabilities of LTE, aiming to improve speed, efficiency, and functionality for future cellular networks​.", "38 series")
        ]
        file_path = 'series_description.json'
        if os.path.isfile(file_path):
            # File exists, read the file and load the JSON content into series_dict
            with open(file_path, 'r') as file:
                series_dict = json.load(file)
        else:
            series_dict = {}
            for desc, series_index in topics_with_series:
                series_dict[series_index] = {}
                series_dict[series_index]["description"]= desc
                series_dict[series_index]["embeddings"]= Query.get_embeddings(desc)
            # File does not exist, write the series_dict to the file as JSON
            with open(file_path, 'w') as file:
                json.dump(series_dict, file, indent=4)

        similarity_coloumn = []
        for embeddings in embeddings_list:
            coef = []
            for series_id in series_dict:
                coef.append(Query.inner_product(embeddings, series_dict[series_id]['embeddings']))
            similarity_coloumn.append(coef)
        return similarity_coloumn

    def preprocessing_softmax(embeddings_list):
        embeddings = np.array(embeddings_list)
        similarity = np.array(Query.get_col2(embeddings))

        X_train_1_tensor = torch.tensor(embeddings, dtype=torch.float32)

        _similarity = torch.from_numpy(similarity)
        _similarity = torch.tensor(similarity, dtype=torch.float32)
        X_train_2_tensor= torch.nn.functional.softmax(10*_similarity, dim=-1)


        dataset = TensorDataset(X_train_1_tensor, X_train_2_tensor)
        gen = torch.Generator(device=device)
        dataloader = DataLoader(dataset, batch_size=128, shuffle=True, generator=gen)

        return dataloader

    def get_embeddings(text):
        # Initialize the OpenAI client
        client = OpenAI()
        # Request embeddings for the list of texts using a different, larger model
        response = client.embeddings.create(
                    input=text,
                    model="text-embedding-3-large",
                    dimensions=1024,
                )
        return response.data[0].embedding

    def predict_wg(self):
        model = NNRouter()
        #model.load_state_dict(torch.load('/content/router_new.pth', map_location=device))
        if not zipfile.is_zipfile('/content/router_new.pth'):
            raise ValueError("The file '/content/router_new.pth' is not a valid zip archive.")

        # Load the state dictionary using a different method
        model = NNRouter()
        model.load_state_dict(torch.load(open('/content/router_new.pth', 'rb'), map_location=device))
        model.to(device)
        model.eval()
        text_list = []
        text_embeddings = Query.get_embeddings_list([self.enhanced_query])
        label_list = []
        embeddings = text_embeddings[self.enhanced_query]
        test_dataloader= Query.preprocessing_softmax([embeddings])
        with torch.no_grad():
            for X1, X2 in test_dataloader:
                # Move data to the same device as the model
                X1, X2 = X1.to(device), X2.to(device)
                original_labels_mapping = np.arange(21, 39)
                outputs = model(X1, X2)
                top_values, top_indices = outputs.topk(5, dim=1)
                # Convert the indices to a numpy array
                predicted_indices = top_indices.cpu().numpy()
                predicted_labels = original_labels_mapping[predicted_indices]
                label_list=predicted_labels
        self.wg = label_list[0]
        print(self.wg)

    def get_question_context_faiss(self, batch, k, use_context=False):
        try:
            faiss_index, faiss_index_to_data_mapping, source_mapping = get_faiss_batch_index(batch)
            if use_context:
                result = find_nearest_neighbors_faiss(self.query, faiss_index, faiss_index_to_data_mapping, k, source_mapping= source_mapping, context=self.context)
            else:
                result = find_nearest_neighbors_faiss(self.query, faiss_index, faiss_index_to_data_mapping, k, source_mapping= source_mapping)

            if isinstance(result, list):
                    self.context = []
                    self.context_source = []
                    for i in range(len(result)):
                        index, data, source = result[i]
                        self.context.append(f"\nRetrieval {i+1}:\n...{data}...\nThis retrieval is performed from the document {source}.\n")
                        self.context_source.append(f"Index: {index}, Source: {source}")
            else:
                self.context = result
        except Exception as e:
            print(f"An error occurred while getting question context: {e}")
            print(traceback.format_exc())
            self.context = "Error in processing"


# Device configuration
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
print(f"Using device: {device}")

# Example usage
query_instance = Query("How does 5G impact telecommunications?", "5G is the fifth generation of cellular networks.")
query_instance.predict_wg()
def generate(question):
        # Constructing the content context from the question object
        content = '\n'.join(question.context)
        prompt = f"""
Please answer to the following question:
{question.query}

Considering the following context:
{content}

Please answer to the following question:
{question.question}
        """

        logging.info("Generated system prompt for OpenAI completion.")

        # Extracting and cleaning the model's response
        predicted_answers_str = phi2(prompt)
        logging.info("Model response generated successfully.")

        context = f"The retrieved context provided to the LLM is:\n{content}"
        return predicted_answers_str, context, question.question
def find_first_number(string):
    """
    Finds the first sequence of digits in a given string and returns it as an integer.

    Parameters:
    - string: The string to search for numbers.

    Returns:
    - An integer representing the first sequence of digits found in the string.
    - Returns None if no digits are found.
    """
    try:
        # Using regular expression to search for the first occurrence of one or more digits
        match = re.search(r'\d+', string)
        if match:
            return int(match.group())  # Convert the found digits into an integer
        else:
            return None  # Return None if no digits are found
    except Exception as e:
        print(f"An error occurred while trying to find a number in the string: {e}")
        return None

def find_option_number(text):
    """
    Finds all the occurrences of numbers preceded by the word 'option' in a given text.

    Parameters:
    - text: The text to search for 'option' followed by numbers.

    Returns:
    - A list of strings, each representing a number found after 'option'. The numbers are returned as strings.
    - If no matches are found, an empty list is returned.
    """
    try:
        text =  text.lower()
        # Define a regular expression pattern to find 'option' followed by non-digit characters (\D*),
        # and then one or more digits (\d+)
        pattern = r'option\D*(\d+)'
        # Find all matches of the pattern in the text
        matches = re.findall(pattern, text)
        return matches  # Return the list of found numbers as strings
    except Exception as e:
        print(f"An error occurred while trying to find option numbers in the text: {e}")
        return []

def check_question_nonjson(question, answer, options):
    """
    This function checks if the answer provided for a non-JSON formatted question is correct.
    It dynamically selects the model based on the model_name provided and constructs a prompt
    for the AI model to generate an answer. It then compares the generated answer with the provided
    answer to determine correctness.

    Parameters:
    - question: A dictionary containing the question, options, and context.
    - model_name: Optional; specifies the model to use. Defaults to 'mistralai/Mixtral-8x7B-Instruct-v0.1'
    if not provided or if the default model is indicated.

    Returns:
    - A tuple containing the updated question dictionary and a boolean indicating correctness.
    """
    # Extracting options from the question dictionary.
    options_text = '\n'.join(options)

    content = '\n'.join(question.context)
    # Constructing the system prompt for the AI model.
    syst_prompt = f"""
    Instruct:
    Please provide the answers to the following multiple choice question.
    {question.query}

    Considering the following context:
    {content}

    Please provide the answers to the following multiple choice question.
    {question.question}
    The output should be in the format: Option

    Options:
    Write only the option number corresponding to the correct answer:\n{options_text}

    Output:
    """
    predicted_answers_str = phi2(syst_prompt)
    predicted_answers_str = predicted_answers_str.split("Output:")[-1]
    print(predicted_answers_str)
    print("Correct answer")
    print(answer)

    # Finding and comparing the predicted answer to the actual answer.
    answer_id = find_option_number(predicted_answers_str)
    print(f"tested option {answer_id}")

    if find_option_number(answer) == answer_id:
        print("Correct\n")
        return question, True
    else:
        print("Wrong\n")
        return question, False
def TelcoRAG(query, answer, options,  api_key= api_key):
    try:
        os.environ["OPENAI_API_KEY"] = api_key
        os.environ["KMP_DUPLICATE_LIB_OK"] = 'TRUE'
        question = Query(query, [])

        question.def_TA_question()

        question.predict_wg()

        document_ds = get_documents(question.wg)

        print(len(document_ds))
        # Chunk documents based on provided chunk size and overlap
        Document_ds = [chunk_doc(doc) for doc in document_ds]


        series_doc = {'Summaries':[]}
        for series_number in question.wg:
            series_doc[f'Series{series_number}'] = []
            for doc in Document_ds:
                if doc[0]['source'][:2].isnumeric():
                    if int(doc[0]['source'][:2]) == series_number:
                        series_doc[f'Series{series_number}'].append(doc)
                else:
                    if doc not in series_doc['Summaries']:
                        series_doc['Summaries'].append(doc)

        series_docs = get_embeddings(series_doc)

        embedded_docs = []
        for serie in series_docs.values():
            embedded_docs.extend([serie])
        question.get_question_context_faiss(batch=embedded_docs, k=10, use_context=False)

        question.candidate_answers()

        old_list =  question.wg
        question.predict_wg()
        new_series = {}
        for series_number in question.wg:
            if series_number not in old_list:
                new_series[f'Series{series_number}'] = []
                for doc in Document_ds:
                    if doc[0]['source'][:2].isnumeric():
                        if int(doc[0]['source'][:2]) == series_number:
                            new_series[f'Series{series_number}'].append(doc)
        new_series = get_embeddings(new_series)
        old_series={'Summaries': series_docs['Summaries']}
        for series_number in question.wg:
            if series_number in old_list:
              old_series[f'Series{series_number}'] = series_docs[f'Series{series_number}']
        embedded_docs = []
        for serie in new_series.values():
            embedded_docs.extend([serie])
        for serie in old_series.values():
            embedded_docs.extend([serie])
        question.get_question_context_faiss(batch=embedded_docs, k=10, use_context=True)

        print(check_question_nonjson(question, answer , options))
        # return response, context, query

    except Exception as e:
        print(f"An error occurred: {e}")
        print(traceback.format_exc())

def choose_random_question(data):
    while True:
        random_question = random.choice(list(data.values()))

        if '3GPP' in random_question['question']:
            return random_question['question']
        else:
            continue





'''
import numpy as np
import pandas as pd
import torch
from string import Template
import json
from tqdm.auto import tqdm
import warnings
from transformers import pipeline

warnings.simplefilter("ignore")

# Check for CUDA availability
device = "cuda:0" if torch.cuda.is_available() else "cpu"

# Load test data
with open("/content/TeleQnA_testing1.txt", "r") as f:
    test_data = f.read()
test_data = json.loads(test_data)

# Function to generate responses and store in context.txt
def generate_responses():
    # Initialize pipeline
    phi3_telecom = pipeline("text-generation", model="EleutherAI/gpt-neo-125M")
    with open("context.txt", "w") as file:
        for question_id, question_details in tqdm(test_data.items()):
            original_query = question_details["question"]
            response = phi3_telecom(original_query, max_new_tokens=50)[0]['generated_text'].strip()
            file.write(f"{question_id}: {response}\n")

# Generate responses and store in context.txt
generate_responses()
'''
